#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
FIXED: Script để extract ĐẦY ĐỦ thông tin từ 207 file .doc thủ tục hành chính
Output: JSON files với cấu trúc chuẩn hóa - ĐẦY ĐỦ 20 TRƯỜNG

FIXES:
1. ✅ Bỏ check "Bước " để capture đầy đủ "Trình tự thực hiện"
2. ✅ Thêm column "mô_tả" vào bảng hinh_thuc_nop (4 columns thay vì 3)
3. ✅ Thêm extraction cho "Từ khóa", "Mô tả", "Cơ quan được ủy quyền"
4. ✅ Thêm columns cho bảng can_cu_phap_ly (ngày_ban_hành, cơ_quan_ban_hành)
"""

import sys
import os
import json
from pathlib import Path
from docx import Document
import glob
from typing import Dict, List, Optional
import re

# Đảm bảo output UTF-8
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')

# Danh sách ĐẦY ĐỦ 20 trường quan trọng
FIELDS_TO_TRACK = [
    "Mã thủ tục",
    "Số quyết định",
    "Tên thủ tục",
    "Cấp thực hiện",
    "Loại thủ tục",
    "Lĩnh vực",
    "Trình tự thực hiện",
    "Cách thức thực hiện",
    "Thành phần hồ sơ",
    "Đối tượng thực hiện",
    "Cơ quan thực hiện",
    "Cơ quan có thẩm quyền",
    "Địa chỉ tiếp nhận HS",
    "Cơ quan được ủy quyền",
    "Cơ quan phối hợp",
    "Kết quả thực hiện",
    "Căn cứ pháp lý",
    "Yêu cầu, điều kiện thực hiện",
    "Từ khóa",
    "Mô tả"
]


def extract_field_value(paragraphs: List, field_name: str) -> str:
    """
    Trích xuất giá trị của một trường từ danh sách paragraphs

    FIXED: Bỏ check "Bước " để không mất nội dung "Trình tự thực hiện"
    """
    field_value = ""
    capturing = False

    for para in paragraphs:
        text = para.text.strip()

        # Bắt đầu capture khi tìm thấy tên trường
        if text.startswith(field_name + ":"):
            # Lấy phần sau dấu ':'
            field_value = text[len(field_name) + 1:].strip()
            capturing = True
            continue

        # Nếu đang capture và gặp trường tiếp theo, dừng lại
        if capturing:
            # Kiểm tra xem có phải là một trường mới không
            is_new_field = any(text.startswith(f + ":") for f in FIELDS_TO_TRACK)

            # FIXED: BỎ CHECK "Bước " - cho phép capture "Bước 1:", "Bước 2:" trong "Trình tự thực hiện"
            # OLD: if is_new_field or text.startswith("Bước "):
            # NEW: chỉ check new field
            if is_new_field:
                break

            # Nếu không phải trường mới, thêm vào giá trị
            if text:
                field_value += " " + text

    return field_value.strip()


def extract_table_data(doc: Document) -> Dict[str, List]:
    """
    Trích xuất dữ liệu từ các bảng trong document
    Tự động nhận biết bảng dựa vào header

    FIXED:
    - Thêm column "mô_tả" cho bảng hinh_thuc_nop (4 columns)
    - Thêm columns "ngày_ban_hành", "cơ_quan_ban_hành" cho bảng can_cu_phap_ly
    """
    table_data = {
        'hinh_thuc_nop': [],
        'thanh_phan_ho_so': [],
        'can_cu_phap_ly': []
    }

    if len(doc.tables) < 1:
        return table_data

    # Duyệt qua tất cả các bảng
    for table in doc.tables:
        if len(table.rows) < 1:
            continue

        # Đọc header (hàng đầu tiên)
        header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]

        # Nhận diện Bảng 1: Hình thức nộp (của "Cách thức thực hiện")
        # FIXED: Thêm column "mô_tả" (4 columns thay vì 3)
        if any('hình thức' in h for h in header_cells) and any('thời hạn' in h for h in header_cells):
            for i, row in enumerate(table.rows):
                if i == 0:  # Bỏ qua header
                    continue
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 3 and cells[0]:
                    # FIXED: Thêm column thứ 4 - mô_tả
                    table_data['hinh_thuc_nop'].append({
                        "hinh_thuc": cells[0],
                        "thoi_han_giai_quyet": cells[1],
                        "phi_le_phi": cells[2],
                        "mo_ta": cells[3] if len(cells) > 3 else ""  # NEW: column mô_tả
                    })

        # Nhận diện Bảng 2: Thành phần hồ sơ
        elif any('tên giấy tờ' in h for h in header_cells) or any('giấy tờ' in h for h in header_cells):
            for i, row in enumerate(table.rows):
                if i == 0:
                    continue
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 1 and cells[0]:
                    # Lấy thêm số lượng và ghi chú nếu có
                    ten_giay_to = cells[0]
                    so_luong = cells[1] if len(cells) > 1 else ""
                    ghi_chu = cells[2] if len(cells) > 2 else ""

                    table_data['thanh_phan_ho_so'].append({
                        "ten_giay_to": ten_giay_to,
                        "so_luong": so_luong,
                        "ghi_chu": ghi_chu
                    })

        # Nhận diện Bảng 3: Căn cứ pháp lý
        # FIXED: Thêm columns "ngày_ban_hành", "cơ_quan_ban_hành"
        elif any('trích yếu' in h for h in header_cells) or any('số ký hiệu' in h for h in header_cells):
            # Tìm vị trí các cột
            so_ky_hieu_col = -1
            trich_yeu_col = -1
            ngay_ban_hanh_col = -1
            co_quan_ban_hanh_col = -1

            for idx, h in enumerate(header_cells):
                if 'số' in h and 'ký hiệu' in h:
                    so_ky_hieu_col = idx
                elif 'trích yếu' in h:
                    trich_yeu_col = idx
                elif 'ngày' in h and 'ban hành' in h:
                    ngay_ban_hanh_col = idx
                elif 'cơ quan' in h and 'ban hành' in h:
                    co_quan_ban_hanh_col = idx

            for i, row in enumerate(table.rows):
                if i == 0:
                    continue
                cells = [cell.text.strip() for cell in row.cells]

                if len(cells) > 0:
                    # Extract các columns
                    so_ky_hieu = cells[so_ky_hieu_col] if so_ky_hieu_col >= 0 and so_ky_hieu_col < len(cells) else (cells[0] if len(cells) > 0 else "")
                    trich_yeu = cells[trich_yeu_col] if trich_yeu_col >= 0 and trich_yeu_col < len(cells) else (cells[1] if len(cells) > 1 else "")
                    ngay_ban_hanh = cells[ngay_ban_hanh_col] if ngay_ban_hanh_col >= 0 and ngay_ban_hanh_col < len(cells) else ""
                    co_quan_ban_hanh = cells[co_quan_ban_hanh_col] if co_quan_ban_hanh_col >= 0 and co_quan_ban_hanh_col < len(cells) else ""

                    if so_ky_hieu or trich_yeu:
                        legal_entry = {
                            "so_ky_hieu": so_ky_hieu,
                            "trich_yeu": trich_yeu
                        }
                        # FIXED: Thêm các columns mới nếu có
                        if ngay_ban_hanh:
                            legal_entry["ngay_ban_hanh"] = ngay_ban_hanh
                        if co_quan_ban_hanh:
                            legal_entry["co_quan_ban_hanh"] = co_quan_ban_hanh

                        table_data['can_cu_phap_ly'].append(legal_entry)

    return table_data


def extract_thu_tuc_id_from_filename(filename: str) -> str:
    """
    Extract ID từ tên file: ChiTietTTHC_1.013124.doc -> 1.013124
    """
    match = re.search(r'_(\d+\.\d+)\.doc', filename)
    if match:
        return match.group(1)
    return filename.replace('ChiTietTTHC_', '').replace('.doc', '')


def analyze_doc_file(file_path: str) -> Optional[Dict]:
    """
    Phân tích một file .doc và trả về cấu trúc JSON chuẩn

    FIXED: Thêm extraction cho TẤT CẢ 20 trường
    """
    try:
        doc = Document(file_path)
        paragraphs = doc.paragraphs
        filename = os.path.basename(file_path)

        # Extract ID từ filename
        thu_tuc_id = extract_thu_tuc_id_from_filename(filename)

        # Trích xuất dữ liệu từ bảng
        table_data = extract_table_data(doc)

        # Tạo metadata (6 trường)
        metadata = {}
        for field in ["Mã thủ tục", "Tên thủ tục", "Số quyết định",
                     "Cấp thực hiện", "Loại thủ tục", "Lĩnh vực"]:
            value = extract_field_value(paragraphs, field)
            # Normalize field name cho JSON key
            key = field.lower().replace(" ", "_").replace(",", "")
            metadata[key] = value

        # Tạo content (11 trường)
        # FIXED: Thêm "Cơ quan được ủy quyền", "Từ khóa", "Mô tả"
        content = {}
        for field in ["Trình tự thực hiện", "Cách thức thực hiện",
                     "Đối tượng thực hiện", "Cơ quan thực hiện",
                     "Cơ quan có thẩm quyền", "Cơ quan được ủy quyền",  # FIXED: Added
                     "Cơ quan phối hợp", "Địa chỉ tiếp nhận HS",
                     "Kết quả thực hiện", "Yêu cầu, điều kiện thực hiện",
                     "Từ khóa", "Mô tả"]:  # FIXED: Added
            value = extract_field_value(paragraphs, field)
            key = field.lower().replace(" ", "_").replace(",", "")
            content[key] = value

        # Tạo JSON structure
        result = {
            "thu_tuc_id": thu_tuc_id,
            "source_file": filename,
            "metadata": metadata,
            "content": content,
            "tables": table_data  # FIXED: Structure already correct from extract_table_data
        }

        return result

    except Exception as e:
        print(f"\n❌ Lỗi khi đọc file {file_path}: {str(e)}")
        import traceback
        traceback.print_exc()
        return None


def main():
    """
    Hàm chính để extract tất cả các file
    """
    print("=" * 80)
    print("FIXED EXTRACTION: ĐẦY ĐỦ 20 TRƯỜNG + TABLE COLUMNS")
    print("=" * 80)
    print()
    print("Fixes:")
    print("  ✅ Bỏ check 'Bước ' để capture đầy đủ 'Trình tự thực hiện'")
    print("  ✅ Thêm column 'mô_tả' vào bảng hinh_thuc_nop (4 columns)")
    print("  ✅ Thêm extraction cho 'Từ khóa', 'Mô tả', 'Cơ quan được ủy quyền'")
    print("  ✅ Thêm columns cho bảng can_cu_phap_ly (ngày, cơ quan ban hành)")
    print()

    # Tìm đường dẫn đến thư mục chứa file .doc
    # Script này nằm trong thu_tuc_rag/src/extraction/
    script_dir = Path(__file__).parent
    project_root = script_dir.parent.parent
    data_dir = project_root / "data"

    # Tìm file .doc ở thư mục gốc (nơi có 207 files)
    root_dir = project_root.parent
    file_pattern = str(root_dir / "ChiTietTTHC_*.doc")

    files = glob.glob(file_pattern)

    # Loại bỏ các file tạm
    files = [f for f in files if not os.path.basename(f).startswith('~$')]

    print(f"📁 Tìm thấy {len(files)} file thủ tục hành chính")
    print(f"📂 Output directory: {data_dir / 'extracted_fixed'}")
    print()

    if len(files) == 0:
        print("⚠️  Không tìm thấy file nào!")
        print(f"   Đã tìm ở: {file_pattern}")
        return

    # Tạo thư mục output (mới để không ghi đè)
    output_dir = data_dir / "extracted_fixed"
    output_dir.mkdir(parents=True, exist_ok=True)

    # Extract tất cả các file
    success_count = 0
    failed_files = []

    for i, file_path in enumerate(files, 1):
        filename = os.path.basename(file_path)
        print(f"\r⏳ Đang xử lý: {i}/{len(files)} - {filename[:50]}", end='', flush=True)

        data = analyze_doc_file(file_path)

        if data:
            # Lưu ra file JSON
            thu_tuc_id = data['thu_tuc_id']
            output_file = output_dir / f"{thu_tuc_id}.json"

            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)

            success_count += 1
        else:
            failed_files.append(filename)

    print("\n\n" + "=" * 80)
    print("KẾT QUẢ EXTRACTION (FIXED)")
    print("=" * 80)
    print(f"✅ Thành công: {success_count}/{len(files)} files")
    print(f"❌ Thất bại: {len(failed_files)} files")

    if failed_files:
        print("\nDanh sách files thất bại:")
        for f in failed_files[:10]:  # Hiển thị max 10 files
            print(f"  - {f}")
        if len(failed_files) > 10:
            print(f"  ... và {len(failed_files) - 10} files khác")

    print(f"\n📊 Dữ liệu đã được lưu tại: {output_dir}")

    # Verify một file mẫu
    if success_count > 0:
        print("\n" + "=" * 80)
        print("VERIFICATION: Kiểm tra file mẫu 1.013124.json")
        print("=" * 80)

        sample_file = output_dir / "1.013124.json"
        if sample_file.exists():
            with open(sample_file, 'r', encoding='utf-8') as f:
                sample_data = json.load(f)

            # Check trình_tự_thực_hiện
            trinh_tu = sample_data["content"].get("trình_tự_thực_hiện", "")
            print(f"\n✅ Trình tự thực hiện: {len(trinh_tu)} chars")
            if trinh_tu and len(trinh_tu) > 100:
                print(f"   Preview: {trinh_tu[:150]}...")
            elif not trinh_tu or trinh_tu == "":
                print(f"   ⚠️  VẪN EMPTY! Cần kiểm tra file .doc gốc")

            # Check hinh_thuc_nop table
            hinh_thuc_nop = sample_data["tables"].get("hinh_thuc_nop", [])
            print(f"\n✅ Bảng hinh_thuc_nop: {len(hinh_thuc_nop)} rows")
            if hinh_thuc_nop and len(hinh_thuc_nop) > 0:
                first_row = hinh_thuc_nop[0]
                print(f"   Columns: {list(first_row.keys())}")
                if "mo_ta" in first_row:
                    print(f"   ✅ Column 'mo_ta' có mặt!")
                else:
                    print(f"   ❌ Column 'mo_ta' THIẾU!")

            # Check từ khóa và mô tả
            tu_khoa = sample_data["content"].get("từ_khóa", "")
            mo_ta = sample_data["content"].get("mô_tả", "")
            print(f"\n✅ Từ khóa: '{tu_khoa}'")
            print(f"✅ Mô tả: '{mo_ta}'")

    print("\n" + "=" * 80)


if __name__ == "__main__":
    main()
