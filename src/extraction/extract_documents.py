#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Script để extract thông tin từ 207 file .doc thủ tục hành chính
Output: JSON files với cấu trúc chuẩn hóa
"""

import sys
import os
import json
from pathlib import Path
from docx import Document
import glob
from typing import Dict, List, Optional
import re

# Đảm bảo output UTF-8
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')

# Danh sách 20 trường quan trọng
FIELDS_TO_TRACK = [
    "Mã thủ tục",
    "Số quyết định",
    "Tên thủ tục",
    "Cấp thực hiện",
    "Loại thủ tục",
    "Lĩnh vực",
    "Trình tự thực hiện",
    "Cách thức thực hiện",
    "Thành phần hồ sơ",
    "Đối tượng thực hiện",
    "Cơ quan thực hiện",
    "Cơ quan có thẩm quyền",
    "Địa chỉ tiếp nhận HS",
    "Cơ quan được ủy quyền",
    "Cơ quan phối hợp",
    "Kết quả thực hiện",
    "Căn cứ pháp lý",
    "Yêu cầu, điều kiện thực hiện",
    "Từ khóa",
    "Mô tả"
]


def extract_field_value(paragraphs: List, field_name: str) -> str:
    """
    Trích xuất giá trị của một trường từ danh sách paragraphs
    """
    field_value = ""
    capturing = False

    for para in paragraphs:
        text = para.text.strip()

        # Bắt đầu capture khi tìm thấy tên trường
        if text.startswith(field_name + ":"):
            # Lấy phần sau dấu ':'
            field_value = text[len(field_name) + 1:].strip()
            capturing = True
            continue

        # Nếu đang capture và gặp trường tiếp theo, dừng lại
        if capturing:
            # Kiểm tra xem có phải là một trường mới không
            is_new_field = any(text.startswith(f + ":") for f in FIELDS_TO_TRACK)
            # Hoặc là dòng "Bước X:"
            if is_new_field or text.startswith("Bước "):
                break
            # Nếu không phải trường mới, thêm vào giá trị
            if text:
                field_value += " " + text

    return field_value.strip()


def extract_table_data(doc: Document) -> Dict[str, List]:
    """
    Trích xuất dữ liệu từ các bảng trong document
    Tự động nhận biết bảng dựa vào header
    """
    table_data = {
        'hinh_thuc_nop': [],
        'thoi_han_giai_quyet': [],
        'phi_le_phi': [],
        'thanh_phan_ho_so': [],
        'can_cu_phap_ly': []
    }

    if len(doc.tables) < 1:
        return table_data

    # Duyệt qua tất cả các bảng
    for table in doc.tables:
        if len(table.rows) < 1:
            continue

        # Đọc header (hàng đầu tiên)
        header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]

        # Nhận diện Bảng 1: Hình thức nộp, Thời hạn, Phí lệ phí
        if any('hình thức' in h for h in header_cells) and any('thời hạn' in h for h in header_cells):
            for i, row in enumerate(table.rows):
                if i == 0:  # Bỏ qua header
                    continue
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 3 and cells[0]:
                    table_data['hinh_thuc_nop'].append(cells[0])
                    table_data['thoi_han_giai_quyet'].append(cells[1])
                    table_data['phi_le_phi'].append(cells[2])

        # Nhận diện Bảng 2: Thành phần hồ sơ
        elif any('tên giấy tờ' in h for h in header_cells) or any('giấy tờ' in h for h in header_cells):
            for i, row in enumerate(table.rows):
                if i == 0:
                    continue
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 1 and cells[0]:
                    # Lấy thêm số lượng và ghi chú nếu có
                    ten_giay_to = cells[0]
                    so_luong = cells[1] if len(cells) > 1 else ""
                    ghi_chu = cells[2] if len(cells) > 2 else ""

                    table_data['thanh_phan_ho_so'].append({
                        "ten_giay_to": ten_giay_to,
                        "so_luong": so_luong,
                        "ghi_chu": ghi_chu
                    })

        # Nhận diện Bảng 3: Căn cứ pháp lý
        elif any('trích yếu' in h for h in header_cells) or any('số ký hiệu' in h for h in header_cells):
            for i, row in enumerate(table.rows):
                if i == 0:
                    continue
                cells = [cell.text.strip() for cell in row.cells]

                # Tìm cột "Trích yếu"
                trich_yeu_col = -1
                so_ky_hieu_col = -1

                for idx, h in enumerate(header_cells):
                    if 'trích yếu' in h:
                        trich_yeu_col = idx
                    if 'số' in h and 'ký hiệu' in h:
                        so_ky_hieu_col = idx

                if len(cells) > max(trich_yeu_col, so_ky_hieu_col):
                    so_ky_hieu = cells[so_ky_hieu_col] if so_ky_hieu_col >= 0 else cells[0]
                    trich_yeu = cells[trich_yeu_col] if trich_yeu_col >= 0 else cells[1] if len(cells) > 1 else ""

                    if so_ky_hieu or trich_yeu:
                        table_data['can_cu_phap_ly'].append({
                            "so_ky_hieu": so_ky_hieu,
                            "trich_yeu": trich_yeu
                        })

    return table_data


def extract_thu_tuc_id_from_filename(filename: str) -> str:
    """
    Extract ID từ tên file: ChiTietTTHC_1.013124.doc -> 1.013124
    """
    match = re.search(r'_(\d+\.\d+)\.doc', filename)
    if match:
        return match.group(1)
    return filename.replace('ChiTietTTHC_', '').replace('.doc', '')


def analyze_doc_file(file_path: str) -> Optional[Dict]:
    """
    Phân tích một file .doc và trả về cấu trúc JSON chuẩn
    """
    try:
        doc = Document(file_path)
        paragraphs = doc.paragraphs
        filename = os.path.basename(file_path)

        # Extract ID từ filename
        thu_tuc_id = extract_thu_tuc_id_from_filename(filename)

        # Trích xuất dữ liệu từ bảng
        table_data = extract_table_data(doc)

        # Tạo metadata
        metadata = {}
        for field in ["Mã thủ tục", "Tên thủ tục", "Số quyết định",
                     "Cấp thực hiện", "Loại thủ tục", "Lĩnh vực"]:
            value = extract_field_value(paragraphs, field)
            # Normalize field name cho JSON key
            key = field.lower().replace(" ", "_").replace(",", "")
            metadata[key] = value

        # Tạo content
        content = {}
        for field in ["Trình tự thực hiện", "Cách thức thực hiện",
                     "Đối tượng thực hiện", "Cơ quan thực hiện",
                     "Cơ quan có thẩm quyền", "Cơ quan phối hợp",
                     "Địa chỉ tiếp nhận HS", "Kết quả thực hiện",
                     "Yêu cầu, điều kiện thực hiện"]:
            value = extract_field_value(paragraphs, field)
            key = field.lower().replace(" ", "_").replace(",", "")
            content[key] = value

        # Tạo JSON structure
        result = {
            "thu_tuc_id": thu_tuc_id,
            "source_file": filename,
            "metadata": metadata,
            "content": content,
            "tables": {
                "hinh_thuc_nop": [
                    {
                        "hinh_thuc": table_data['hinh_thuc_nop'][i] if i < len(table_data['hinh_thuc_nop']) else "",
                        "thoi_han_giai_quyet": table_data['thoi_han_giai_quyet'][i] if i < len(table_data['thoi_han_giai_quyet']) else "",
                        "phi_le_phi": table_data['phi_le_phi'][i] if i < len(table_data['phi_le_phi']) else ""
                    }
                    for i in range(max(len(table_data['hinh_thuc_nop']),
                                      len(table_data['thoi_han_giai_quyet']),
                                      len(table_data['phi_le_phi'])))
                ] if table_data['hinh_thuc_nop'] or table_data['thoi_han_giai_quyet'] or table_data['phi_le_phi'] else [],
                "thanh_phan_ho_so": table_data['thanh_phan_ho_so'],
                "can_cu_phap_ly": table_data['can_cu_phap_ly']
            }
        }

        return result

    except Exception as e:
        print(f"❌ Lỗi khi đọc file {file_path}: {str(e)}")
        return None


def main():
    """
    Hàm chính để extract tất cả các file
    """
    print("=" * 80)
    print("EXTRACT DỮ LIỆU TỪ 207 FILE THỦ TỤC HÀNH CHÍNH")
    print("=" * 80)
    print()

    # Tìm đường dẫn đến thư mục chứa file .doc
    # Script này nằm trong thu_tuc_rag/src/extraction/
    script_dir = Path(__file__).parent
    project_root = script_dir.parent.parent
    data_dir = project_root / "data"

    # Tìm file .doc ở thư mục gốc (nơi có 207 files)
    root_dir = project_root.parent
    file_pattern = str(root_dir / "ChiTietTTHC_*.doc")

    files = glob.glob(file_pattern)

    # Loại bỏ các file tạm
    files = [f for f in files if not os.path.basename(f).startswith('~$')]

    print(f"📁 Tìm thấy {len(files)} file thủ tục hành chính")
    print(f"📂 Output directory: {data_dir / 'extracted'}")
    print()

    if len(files) == 0:
        print("⚠️  Không tìm thấy file nào!")
        print(f"   Đã tìm ở: {file_pattern}")
        return

    # Tạo thư mục output
    output_dir = data_dir / "extracted"
    output_dir.mkdir(parents=True, exist_ok=True)

    # Extract tất cả các file
    success_count = 0
    failed_files = []

    for i, file_path in enumerate(files, 1):
        filename = os.path.basename(file_path)
        print(f"\r⏳ Đang xử lý: {i}/{len(files)} - {filename}", end='', flush=True)

        data = analyze_doc_file(file_path)

        if data:
            # Lưu ra file JSON
            thu_tuc_id = data['thu_tuc_id']
            output_file = output_dir / f"{thu_tuc_id}.json"

            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)

            success_count += 1
        else:
            failed_files.append(filename)

    print("\n\n" + "=" * 80)
    print("KẾT QUẢ EXTRACTION")
    print("=" * 80)
    print(f"✅ Thành công: {success_count}/{len(files)} files")
    print(f"❌ Thất bại: {len(failed_files)} files")

    if failed_files:
        print("\nDanh sách files thất bại:")
        for f in failed_files[:10]:  # Hiển thị max 10 files
            print(f"  - {f}")
        if len(failed_files) > 10:
            print(f"  ... và {len(failed_files) - 10} files khác")

    print(f"\n📊 Dữ liệu đã được lưu tại: {output_dir}")
    print("=" * 80)


if __name__ == "__main__":
    main()
